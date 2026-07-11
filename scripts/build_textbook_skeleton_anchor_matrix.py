from __future__ import annotations

import argparse
import csv
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional runtime dependency
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = Path(r"E:\BigMouse\0.CDSS文献诊疗指南材料PDF\心血管内科\书籍教材\《内科学（第10版）》.docx")
DEFAULT_PDF = Path(r"E:\BigMouse\0.CDSS文献诊疗指南材料PDF\心血管内科\书籍教材\《内科学（第10版）》.pdf")
DEFAULT_P0 = ROOT / "心血管内科文献集合" / "99_教材骨架质量审计_textbook_skeleton_audit" / "20260708_074929" / "priority_four_p0_textbook_skeleton_audit.csv"
DEFAULT_OUT = ROOT / "心血管内科文献集合" / "00_教材骨架库_foundation_skeleton" / "20260708_textbook_anchor_matrix"


CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百]+章\s*(.+)$")
SECTION_RE = re.compile(r"^第[一二三四五六七八九十百]+节\s*\|\s*(.+)$")
ITEM_RE = re.compile(r"^[一二三四五六七八九十百]+\s*[、，]\s*(.+)$")
SLOT_RE = re.compile(r"^【([^】]+)】\s*(.*)$")


SLOT_MAP = {
    "概述": "overview",
    "定义": "overview",
    "血压分类和定义": "overview",
    "类型": "classification_risk",
    "临床分型": "classification_risk",
    "临床分类": "classification_risk",
    "临床表现": "clinical_manifestation",
    "临床特征": "clinical_manifestation",
    "病因": "etiology",
    "病因和危险因素": "etiology",
    "病因和发病机制": "etiology",
    "病因与分类": "etiology",
    "发病机制": "pathogenesis",
    "病理": "pathogenesis",
    "病理改变": "pathogenesis",
    "病理解剖": "pathogenesis",
    "病理解剖和病理生理": "pathogenesis",
    "病理改变和病理生理": "pathogenesis",
    "病理生理": "pathogenesis",
    "辅助检查": "exam_lab",
    "实验室检查": "exam_lab",
    "实验室和辅助检查": "exam_lab",
    "诊断": "diagnosis_differential",
    "诊断与鉴别诊断": "diagnosis_differential",
    "鉴别诊断": "diagnosis_differential",
    "分级和危险分层": "classification_risk",
    "分期与分级": "classification_risk",
    "危险评估和预后": "classification_risk",
    "治疗": "treatment",
    "防治": "treatment",
    "临床管理": "treatment",
    "猝死预防": "treatment",
    "预后": "prognosis_followup_prevention",
    "预防": "prognosis_followup_prevention",
}


PRIORITY_CATEGORIES = ("冠状动脉粥样硬化性心脏病", "心肌病", "心力衰竭", "心律失常", "心脏骤停与心脏性猝死")


PDF_WINDOWS = {
    "心力衰竭": (207, 221),
    "心律失常": (221, 258),
    "心脏骤停与心脏性猝死": (221, 258),
    "冠状动脉粥样硬化性心脏病": (258, 306),
    "心肌病": (306, 319),
}


TITLE_HINTS = {
    "DIS-CARD-CAD-CHD": ["冠状动脉粥样硬化性心脏病", "冠心病"],
    "DIS-CARD-CAD-ATHEROSCLEROSIS": ["动脉粥样硬化"],
    "DIS-CARD-CAD-CCS": ["慢性冠状动脉综合征", "慢性冠脉综合征", "CCS"],
    "DIS-CARD-CAD-STABLE-ANGINA": ["稳定型心绞痛", "stable angina"],
    "DIS-CARD-CAD-ICM": ["缺血性心肌病"],
    "DIS-CARD-CAD-SILENT-ISCHEMIA": ["隐匿型冠心病", "无症状性冠心病", "无症状心肌缺血"],
    "DIS-CARD-CAD-ACS": ["急性冠脉综合征", "ACS"],
    "DIS-CARD-CAD-UA": ["不稳定型心绞痛", "UA"],
    "DIS-CARD-CAD-NSTEMI": ["非 ST 段抬高型心肌梗死", "非ST段抬高型心肌梗死", "NSTEMI"],
    "DIS-CARD-CAD-STEMI": ["急性 ST 段抬高型心肌梗死", "ST 段抬高型心肌梗死", "ST段抬高型心肌梗死", "STEMI"],
    "DIS-CARD-CAD-AMI": ["急性心肌梗死", "AMI"],
    "DIS-CARD-CAD-SPASM": ["冠脉痉挛", "冠状动脉痉挛"],
    "DIS-CARD-CAD-OLD-MI": ["陈旧性 MI", "陈旧性心肌梗死", "慢性期改变"],
    "DIS-CARD-CAD-POST-MI-SYNDROME": ["心肌梗死后综合征"],
    "DIS-CARD-HF": ["心力衰竭", "心衰"],
    "DIS-CARD-HF-CHF": ["慢性心力衰竭", "慢性心衰", "CHF"],
    "DIS-CARD-HF-AHF": ["急性心力衰竭", "急性心衰", "AHF"],
    "DIS-CARD-HF-LEFT": ["左心衰竭"],
    "DIS-CARD-HF-RIGHT": ["右心衰竭"],
    "DIS-CARD-HF-BIVENTRICULAR": ["全心衰竭"],
    "DIS-CARD-HF-HFrEF": ["射血分数降低", "HFrEF", "LVEF≤40"],
    "DIS-CARD-HF-HFmrEF": ["射血分数轻度降低", "HFmrEF", "LVEF 为 41%～49%"],
    "DIS-CARD-HF-HFpEF": ["射血分数保留", "HFpEF"],
    "DIS-CARD-HF-POST-MI": ["心肌梗死后心力衰竭", "心肌梗死", "心衰"],
    "DIS-CARD-HF-DIALYSIS-CHF": ["透析患者慢性心力衰竭", "透析"],
    "DIS-CARD-ARR-SINUS": ["窦性心律失常"],
    "DIS-CARD-ARR-SB": ["窦性心动过缓"],
    "DIS-CARD-ARR-SAB": ["窦房传导阻滞"],
    "DIS-CARD-ARR-SND": ["病态窦房结综合征", "窦房结功能障碍"],
    "DIS-CARD-ARR-AF": ["心房颤动", "AF"],
    "DIS-CARD-ARR-AFL": ["心房扑动"],
    "DIS-CARD-ARR-AT": ["房性心动过速"],
    "DIS-CARD-ARR-SVT": ["室上性心动过速", "SVT"],
    "DIS-CARD-ARR-PSVT": ["阵发性室上性心动过速", "PSVT"],
    "DIS-CARD-ARR-AVRT": ["房室折返性心动过速"],
    "DIS-CARD-ARR-AVNRT": ["房室结折返性心动过速"],
    "DIS-CARD-ARR-WPW": ["预激综合征"],
    "DIS-CARD-ARR-VA": ["室性心律失常"],
    "DIS-CARD-ARR-PVC": ["室性期前收缩"],
    "DIS-CARD-ARR-VT": ["室性心动过速"],
    "DIS-CARD-ARR-NSVT": ["非持续性室性心动过速"],
    "DIS-CARD-ARR-TDP": ["尖端扭转型室性心动过速"],
    "DIS-CARD-ARR-VF": ["心室扑动", "心室颤动"],
    "DIS-CARD-ARR-AVB": ["房室传导阻滞"],
    "DIS-CARD-ARR-AVB1": ["一度房室传导阻滞"],
    "DIS-CARD-ARR-AVB2": ["二度房室传导阻滞"],
    "DIS-CARD-ARR-AVB3": ["三度房室传导阻滞", "完全性房室传导阻滞"],
    "DIS-CARD-ARR-BBB": ["束支传导阻滞", "室内传导阻滞"],
    "DIS-CARD-ARR-BRADY": ["缓慢性心律失常"],
    "DIS-CARD-ARR-LQTS": ["长 QT", "长QT"],
    "DIS-CARD-ARR-SQTS": ["短 QT", "短QT"],
    "DIS-CARD-ARR-BRUGADA": ["Brugada"],
    "DIS-CARD-ARR-CPVT": ["儿茶酚胺敏感性多形性室性心动过速", "CPVT"],
    "DIS-CARD-ARR-ERS": ["早期复极"],
    "DIS-CARD-CM-HCM": ["肥厚型心肌病", "HCM"],
    "DIS-CARD-CM-DCM": ["扩张型心肌病", "DCM"],
    "DIS-CARD-CM-NDLVCM": ["非扩张型左心室心肌病", "NDLVC"],
    "DIS-CARD-CM-ARVC": ["致心律失常性右心室心肌病", "ARVC"],
    "DIS-CARD-CM-RCM": ["限制型心肌病", "RCM"],
    "DIS-CARD-CM-MYOCARDITIS": ["心肌炎"],
    "DIS-CARD-CM-ACM": ["致心律失常性心肌病", "ACM"],
    "DIS-CARD-CM-ABVC": ["致心律失常性双心室心肌病"],
    "DIS-CARD-CM-ALVC": ["致心律失常性左心室心肌病"],
    "DIS-CARD-CM-ATRIAL": ["心房心肌病"],
    "DIS-CARD-CM-FABRY": ["Fabry", "法布雷"],
    "DIS-CARD-CM-AMYLOID": ["淀粉样变", "心脏淀粉样变"],
    "DIS-CARD-SCD-SUDDEN": ["心源性猝死", "心脏性猝死", "猝死"],
    "DIS-CARD-SCD-ARREST": ["心脏骤停", "cardiac arrest"],
}


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    text = re.sub(r"N\s*O\s*T\s*E\s*S", "", text, flags=re.IGNORECASE).strip()
    text = text.replace("N OTE S", "").replace("NOTE S", "")
    return text


def compact_preview(text: str, limit: int = 260) -> str:
    text = clean_text(text)
    return text if len(text) <= limit else text[:limit] + "..."


@dataclass
class Paragraph:
    idx: int
    text: str


@dataclass
class Segment:
    level: str
    title: str
    start_idx: int
    end_idx: int
    chapter_title: str
    section_title: str
    paragraphs: list[Paragraph] = field(default_factory=list)
    slots: dict[str, str] = field(default_factory=dict)
    overview: str = ""
    pdf_page: int | None = None

    @property
    def section_path(self) -> str:
        parts = [p for p in [self.chapter_title, self.section_title, self.title if self.title != self.section_title else ""] if p]
        return " > ".join(parts)

    @property
    def text(self) -> str:
        return clean_text(" ".join(p.text for p in self.paragraphs))


def read_docx_paragraphs(path: Path) -> list[Paragraph]:
    doc = Document(str(path))
    out = []
    for idx, para in enumerate(doc.paragraphs):
        text = clean_text(para.text)
        if not text:
            continue
        if text.isdigit() or text in {"N OTE S", "NOTE S", "本章数字资源"}:
            continue
        out.append(Paragraph(idx=idx, text=text))
    return out


def find_structural_segments(paragraphs: list[Paragraph]) -> list[Segment]:
    chapter_title = ""
    section_title = ""
    starts: list[tuple[str, str, int, str, str]] = []
    for p in paragraphs:
        text = p.text
        m = CHAPTER_RE.match(text)
        if m:
            chapter_title = clean_text(m.group(1))
            section_title = ""
            starts.append(("chapter", chapter_title, p.idx, chapter_title, section_title))
            continue
        m = SECTION_RE.match(text)
        if m:
            section_title = clean_text(m.group(1))
            starts.append(("section", section_title, p.idx, chapter_title, section_title))
            continue
        m = ITEM_RE.match(text)
        if m and section_title:
            title = clean_text(m.group(1))
            # 药物/检查等槽位内的小条目很多，只有较像疾病或临床亚型的顶层条目才作为子段。
            if any(key in title for key in ["心绞痛", "心肌病", "冠心病", "心肌梗死", "心动过速", "心动过缓", "期前收缩", "早搏", "心房", "心室", "传导阻滞", "窦房结", "预激", "扑动", "颤动", "猝死", "骤停", "束支", "痉挛", "QT", "Brugada", "早期复极", "儿茶酚胺", "综合征"]):
                starts.append(("item", title, p.idx, chapter_title, section_title))

    # Only cardiovascular chapter window: from 心力衰竭 to before non-cardiology spillover.
    cv_starts = [s for s in starts if s[2] >= 3905 and s[2] < 7600]
    segments: list[Segment] = []
    for i, (level, title, start, ch, sec) in enumerate(cv_starts):
        end = cv_starts[i + 1][2] if i + 1 < len(cv_starts) else 7600
        seg_paras = [p for p in paragraphs if start <= p.idx < end]
        seg = Segment(level=level, title=title, start_idx=start, end_idx=end - 1, chapter_title=ch, section_title=sec, paragraphs=seg_paras)
        if seg.title == "动脉粥样硬化":
            seg.chapter_title = "动脉粥样硬化和冠状动脉粥样硬化性心脏病"
            seg.section_title = "动脉粥样硬化"
        if seg.title == "冠状动脉痉挛":
            seg.chapter_title = "动脉粥样硬化和冠状动脉粥样硬化性心脏病"
            seg.section_title = "冠状动脉疾病的其他表现形式"
        enrich_segment_slots(seg)
        segments.append(seg)
    return segments


def enrich_segment_slots(seg: Segment) -> None:
    current_slot = "overview"
    slot_chunks: dict[str, list[str]] = {"overview": []}
    for p in seg.paragraphs[1:]:
        text = p.text
        # 新小节标题自身不进入 overview。
        if ITEM_RE.match(text) and p.idx == seg.start_idx:
            continue
        m = SLOT_RE.match(text)
        if m:
            slot_name = clean_text(m.group(1))
            current_slot = SLOT_MAP.get(slot_name, slot_name)
            slot_chunks.setdefault(current_slot, [])
            rest = clean_text(m.group(2))
            if rest:
                slot_chunks[current_slot].append(rest)
            continue
        # 仅过滤明显的页眉脚；正文里的（一）属于当前 slot 内容。
        if text in {"N OTE S", "NOTE S"}:
            continue
        slot_chunks.setdefault(current_slot, []).append(text)

    seg.slots = {k: clean_text(" ".join(v)) for k, v in slot_chunks.items() if clean_text(" ".join(v))}
    seg.overview = seg.slots.get("overview", "")


def load_p0_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def hints_for(row: dict[str, str]) -> list[str]:
    hints = [row["disease_name"]]
    hints.extend(TITLE_HINTS.get(row["disease_code"], []))
    out = []
    seen = set()
    for h in hints:
        h = clean_text(h)
        if h and h not in seen:
            out.append(h)
            seen.add(h)
    return out


def safe_title_match(title: str, hint: str) -> tuple[int, str] | None:
    title = clean_text(title)
    hint = clean_text(hint)
    if not title or not hint:
        return None
    if title == hint:
        return 100, "exact_title"
    # 避免“不稳定型心绞痛”误命中“稳定型心绞痛”。
    if title.startswith("稳定型") and hint.startswith("不稳定型"):
        return None
    if title.startswith("窦性心律失常") and hint in {"窦性心动过缓", "窦房传导阻滞", "窦房结功能障碍"}:
        return None
    # 过短英文缩写不能用于标题子串匹配，避免 UA/AF/SB 这类缩写误吸附。
    if re.fullmatch(r"[A-Za-z]{1,3}", hint):
        return None
    if hint in title:
        if ("和" in title or "/" in title or "及" in title) and title != hint:
            return 82, "combined_title"
        return 92, "alias_title"
    # 仅允许较长标题作为疾病全称的父标题；否定词场景已在上面排除。
    if len(title) >= 7 and title in hint:
        return 90, "alias_title"
    return None


def category_window(row: dict[str, str]) -> tuple[int, int] | None:
    category = row.get("category", "")
    for key, window in PDF_WINDOWS.items():
        if key in category:
            return window
    return None


def candidate_segments(row: dict[str, str], segments: list[Segment]) -> list[tuple[int, str, Segment, str]]:
    hints = hints_for(row)
    category = row.get("category", "")
    scored: list[tuple[int, str, Segment, str]] = []
    for seg in segments:
        title_text = clean_text(seg.title)
        strong_title_match = any((m := safe_title_match(title_text, h)) and m[0] >= 100 for h in hints)
        # 大类约束，避免跨章误配。心脏骤停/猝死允许在心律失常章节查找。
        if not strong_title_match and "冠状动脉" in category and "冠状动脉" not in seg.chapter_title and "动脉粥样硬化" not in seg.chapter_title:
            continue
        if not strong_title_match and "心力衰竭" in category and "心力衰竭" not in seg.chapter_title:
            continue
        if not strong_title_match and "心肌病" in category and "心肌疾病" not in seg.chapter_title and "冠状动脉" not in category:
            continue
        if not strong_title_match and "心律失常" in category and "心律失常" not in seg.chapter_title:
            continue
        body_text = seg.text
        for h in hints:
            h_norm = clean_text(h)
            if not h_norm:
                continue
            title_match = safe_title_match(title_text, h_norm)
            if title_match:
                score, match_type = title_match
                scored.append((score, match_type, seg, h_norm))
            elif h_norm in seg.overview:
                scored.append((82, "overview_mention", seg, h_norm))
            elif h_norm in body_text:
                scored.append((62, "body_mention", seg, h_norm))
    scored.sort(key=lambda x: (-x[0], x[2].start_idx))
    return scored


def sentence_with_hint(text: str, hints: Iterable[str]) -> str:
    text = clean_text(text)
    if not text:
        return ""
    parts = [clean_text(p) for p in re.split(r"(?<=[。；;])", text) if clean_text(p)]
    scored: list[tuple[int, int, str]] = []
    strong_definitional_markers = ["是指", "称为", "是一", "是一组", "是一类", "合称", "简称"]
    weak_definitional_markers = ["是", "指", "定义", "包括", "为"]
    bad_markers = ["表 ", "表3", "图 ", "见表", "见图", "诱发因素", "常见于", "治疗", "可发生", "伴有"]
    for h in hints:
        for idx, part in enumerate(parts):
            if not h or h not in part:
                continue
            score = 20
            if any(m in part for m in strong_definitional_markers):
                score += 60
            elif any(m in part for m in weak_definitional_markers):
                score += 30
            if any(m in part for m in bad_markers):
                score -= 40
            if re.fullmatch(r"[A-Za-z]{1,5}", h):
                score += 8
            if len(part) < 20:
                score -= 10
            scored.append((score, -idx, part))
    if scored:
        scored.sort(reverse=True)
        return clean_text(scored[0][2])
    return clean_text(parts[0] if parts else text)


def build_definition(row: dict[str, str], match: tuple[int, str, Segment, str] | None) -> tuple[str, str]:
    if not match:
        return "", ""
    _score, match_type, seg, hit = match
    hints = hints_for(row)
    if match_type in {"exact_title", "alias_title", "overview_mention"} and seg.overview:
        definition = sentence_with_hint(seg.overview, hints) or seg.overview
        description = seg.overview
        return compact_preview(definition, 500), compact_preview(description, 1200)
    # 对左/右/全心衰、EF 分型、传导阻滞分度等，只取命中句作为定义候选。
    definition = sentence_with_hint(seg.text, hints)
    return compact_preview(definition, 500), compact_preview(definition, 1200)


def looks_like_header_only_definition(text: str) -> bool:
    text = clean_text(text)
    if not text:
        return True
    if len(text) < 16:
        return True
    if re.match(r"^第[一二三四五六七八九十百]+[章节]\s*\|?\s*", text) and not any(m in text for m in ["是", "指", "称为"]):
        return True
    return False


def load_pdf_pages(pdf_path: Path, windows: list[tuple[int, int]]) -> dict[int, str]:
    if PdfReader is None or not pdf_path.exists():
        return {}
    wanted: set[int] = set()
    for start, end in windows:
        wanted.update(range(start, end + 1))
    reader = PdfReader(str(pdf_path))
    pages: dict[int, str] = {}
    for page_no in sorted(wanted):
        idx = page_no - 1
        if idx < 0 or idx >= len(reader.pages):
            continue
        try:
            pages[page_no] = clean_text(reader.pages[idx].extract_text() or "")
        except Exception:
            pages[page_no] = ""
    return pages


def find_pdf_page(row: dict[str, str], pdf_pages: dict[int, str]) -> int | None:
    window = category_window(row)
    if not window:
        return None
    start, end = window
    hints = hints_for(row)
    for page_no in range(start, end + 1):
        text = pdf_pages.get(page_no, "")
        if not text:
            continue
        for h in hints:
            if h and h in text:
                return page_no
    return None


def write_csv(path: Path, rows: list[dict[str, object]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def main() -> int:
    ap = argparse.ArgumentParser(description="Build textbook skeleton anchor matrix from Internal Medicine DOCX/PDF.")
    ap.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    ap.add_argument("--pdf", type=Path, default=DEFAULT_PDF)
    ap.add_argument("--p0", type=Path, default=DEFAULT_P0)
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = ap.parse_args()

    out_dir = args.out
    out_dir.mkdir(parents=True, exist_ok=True)
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    paragraphs = read_docx_paragraphs(args.docx)
    segments = find_structural_segments(paragraphs)
    p0_rows = load_p0_rows(args.p0)

    windows = sorted(set(PDF_WINDOWS.values()))
    pdf_pages = load_pdf_pages(args.pdf, windows)

    outline_rows: list[dict[str, object]] = []
    for seg in segments:
        if not any(key in seg.chapter_title for key in ["心力衰竭", "心律失常", "动脉粥样硬化", "冠状动脉", "心肌疾病", "高血压", "心脏瓣膜", "心包疾病", "感染性心内膜炎", "先天性心血管病", "主动脉", "周围血管", "心血管神经症", "肿瘤"]):
            continue
        outline_rows.append({
            "level": seg.level,
            "chapter_title": seg.chapter_title,
            "section_title": seg.section_title,
            "title": seg.title,
            "start_paragraph": seg.start_idx,
            "end_paragraph": seg.end_idx,
            "source_section_path": seg.section_path,
            "has_overview": "yes" if seg.overview else "no",
            "overview_preview": compact_preview(seg.overview),
            "slots": "|".join(sorted(k for k in seg.slots if k != "overview")),
        })

    matrix_rows: list[dict[str, object]] = []
    repair_rows: list[dict[str, object]] = []
    for row in p0_rows:
        matches = candidate_segments(row, segments)
        best = matches[0] if matches else None
        if best:
            score, match_type, seg, hit = best
            definition, description = build_definition(row, best)
            pdf_page = find_pdf_page(row, pdf_pages)
            has_target_overview = bool(seg.overview and match_type in {"exact_title", "alias_title", "overview_mention"})
            status = "ready_for_review" if score >= 82 and definition else "needs_manual_anchor_review"
            if score >= 92 and definition and has_target_overview:
                status = "ready_for_import_after_sampling"
            if score < 82:
                status = "needs_manual_anchor_review"
            if looks_like_header_only_definition(definition):
                status = "needs_manual_anchor_review"
            source_section_path = seg.section_path
            start_para = seg.start_idx
            end_para = seg.end_idx
            slots = "|".join(sorted(k for k in seg.slots if k != "overview"))
        else:
            score = 0
            match_type = "no_textbook_anchor"
            hit = ""
            definition = ""
            description = ""
            pdf_page = None
            status = "needs_guideline_or_manual_source"
            source_section_path = ""
            start_para = ""
            end_para = ""
            slots = ""

        matrix = {
            "generated_at": generated_at,
            "priority_scope_note": "冠心病、心肌病、心力衰竭、心律失常为优先验证集；通过后按同规则自动扩大到心血管内科其他疾病。",
            "category": row["category"],
            "disease_code": row["disease_code"],
            "disease_name": row["disease_name"],
            "match_status": status,
            "match_type": match_type,
            "match_score": score,
            "hit_text": hit,
            "source_section_path": source_section_path,
            "docx_paragraph_start": start_para,
            "docx_paragraph_end": end_para,
            "pdf_page_start": pdf_page or "",
            "pdf_page_end": pdf_page or "",
            "skeleton_slot": "overview" if definition else "",
            "knowledge_layer": "textbook_core" if definition else "",
            "definition_candidate": definition,
            "description_candidate": description,
            "available_slots_in_source": slots,
        }
        matrix_rows.append(matrix)
        if status in {"ready_for_import_after_sampling", "ready_for_review"}:
            repair_rows.append({
                "action": "update_disease_textbook_core_definition",
                "disease_code": row["disease_code"],
                "disease_name": row["disease_name"],
                "definition": definition,
                "description": description,
                "source_type": "authoritative_textbook",
                "source_name": "《内科学（第10版）》",
                "source_section_path": source_section_path,
                "docx_paragraph_start": start_para,
                "docx_paragraph_end": end_para,
                "pdf_page_start": pdf_page,
                "pdf_page_end": pdf_page,
                "skeleton_slot": "overview",
                "knowledge_layer": "textbook_core",
                "clinical_review_status": "pending_textbook_anchor_sampling",
                "import_policy": "must_sample_before_server_write",
            })

    write_csv(
        out_dir / "textbook_cardiology_chapter_outline_20260708.csv",
        outline_rows,
        ["level", "chapter_title", "section_title", "title", "start_paragraph", "end_paragraph", "source_section_path", "has_overview", "overview_preview", "slots"],
    )
    write_csv(
        out_dir / "textbook_skeleton_matrix_priority_four_20260708.csv",
        matrix_rows,
        [
            "generated_at",
            "priority_scope_note",
            "category",
            "disease_code",
            "disease_name",
            "match_status",
            "match_type",
            "match_score",
            "hit_text",
            "source_section_path",
            "docx_paragraph_start",
            "docx_paragraph_end",
            "pdf_page_start",
            "pdf_page_end",
            "skeleton_slot",
            "knowledge_layer",
            "definition_candidate",
            "description_candidate",
            "available_slots_in_source",
        ],
    )
    with (out_dir / "p0_definition_repair_input_priority_four_20260708.jsonl").open("w", encoding="utf-8") as f:
        for item in repair_rows:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")

    summary = {
        "generated_at": generated_at,
        "docx": str(args.docx),
        "pdf": str(args.pdf),
        "p0_input": str(args.p0),
        "segment_count": len(segments),
        "p0_disease_count": len(p0_rows),
        "repair_candidate_count": len(repair_rows),
        "status_counts": {},
        "match_type_counts": {},
        "scope_policy": "四类只是优先验证集；通过硬闸门后自动扩大到心血管内科其他疾病。",
        "outputs": {
            "outline": str(out_dir / "textbook_cardiology_chapter_outline_20260708.csv"),
            "matrix": str(out_dir / "textbook_skeleton_matrix_priority_four_20260708.csv"),
            "repair_jsonl": str(out_dir / "p0_definition_repair_input_priority_four_20260708.jsonl"),
        },
    }
    for row in matrix_rows:
        summary["status_counts"][row["match_status"]] = summary["status_counts"].get(row["match_status"], 0) + 1
        summary["match_type_counts"][row["match_type"]] = summary["match_type_counts"].get(row["match_type"], 0) + 1
    (out_dir / "summary_20260708.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
