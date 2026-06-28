import csv
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.parse_docx_batch import parse_included_docx


class ParseDocxBatchTests(unittest.TestCase):
    def test_extracts_paragraphs_and_tables_with_stable_segments(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            batch = root / "batch"
            manifest_dir = batch / "01_source_manifest"
            manifest_dir.mkdir(parents=True)
            docx_path = root / "心肌病指南.docx"

            document = Document()
            document.add_heading("肥厚型心肌病", level=1)
            document.add_paragraph("诊断与治疗建议")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "指标"
            table.cell(0, 1).text = "阈值"
            table.cell(1, 0).text = "LVEF"
            table.cell(1, 1).text = "50%"
            document.save(docx_path)

            fields = [
                "batch_id",
                "document_id",
                "file_name",
                "full_path",
                "extension",
                "inclusion_status",
                "source_type",
            ]
            with (manifest_dir / "source_documents_manifest.csv").open(
                "w", encoding="utf-8-sig", newline=""
            ) as handle:
                writer = csv.DictWriter(handle, fieldnames=fields)
                writer.writeheader()
                writer.writerow(
                    {
                        "batch_id": "BATCH-TEST",
                        "document_id": "DOC-DOCX",
                        "file_name": docx_path.name,
                        "full_path": str(docx_path),
                        "extension": ".docx",
                        "inclusion_status": "included",
                        "source_type": "guideline",
                    }
                )

            summary = parse_included_docx(batch)
            self.assertEqual(summary["document_count"], 1)
            self.assertGreater(summary["segment_count"], 2)
            clean_text = (batch / "03_clean_text" / "DOC-DOCX.clean.txt").read_text(
                encoding="utf-8-sig"
            )
            self.assertIn("肥厚型心肌病", clean_text)
            self.assertIn("LVEF\t50%", clean_text)
            segments = [
                json.loads(line)
                for line in (batch / "03_clean_text" / "segment_index_docx.jsonl")
                .read_text(encoding="utf-8-sig")
                .splitlines()
            ]
            self.assertTrue(all(row["document_id"] == "DOC-DOCX" for row in segments))


if __name__ == "__main__":
    unittest.main()
